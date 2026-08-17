"""
rag_core.py
===========
This is the ENGINE of the whole project. 1_ingest.py, 2_ask.py, and
app.py (the Streamlit demo) all import functions from here instead
of repeating the same logic three times.

If your TL asks "how does retrieval actually work?" - the answer is
entirely in this one file. Read this file top to bottom and you can
explain the whole system.

THE TWO HALVES OF RAG:
  OFFLINE (run when documents change): load -> chunk -> embed -> store
  ONLINE  (run on every question):     embed question -> compare ->
                                         pick best chunks -> ask LLM

THIS VERSION USES GOOGLE'S GEMINI API instead of OpenAI, because the
Gemini API has a genuinely free tier that needs no credit card - just
a Google account. See README.md section 2 for how to get a key.
Only embed_text() and generate_answer() below talk to Gemini -
everything else (chunking, cosine similarity, retrieval logic) is
plain Python/numpy and would be identical no matter which LLM
provider you used.
"""

import os
import glob
import json
import time
import random
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed
import sqlite3 as _sqlite3_db  # aliased: distinct from chat_store.py's own sqlite3 usage, this one reads DATA files, not app storage
from email import policy as _email_policy
from email.parser import BytesParser as _EmailBytesParser
from html.parser import HTMLParser
import numpy as np
from dotenv import load_dotenv
from google import genai
from google.genai import types

# These three are only needed for PDF/Word/Excel ingestion. Imported
# defensively so that a project with only .txt files still works even
# if these haven't been pip installed yet - see extract_text_from_*()
# below, where a clear, actionable error is raised only when a file
# of that specific type is actually encountered.
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    import openpyxl
except ImportError:
    openpyxl = None

# Only needed for Outlook .msg files specifically - .eml (the
# universal, portable email format every client can export to,
# including Outlook) is handled by Python's own built-in email
# module just above, with no extra package required at all.
try:
    import extract_msg
except ImportError:
    extract_msg = None

# FAISS is the actual similarity-search engine as of this version -
# see the VectorStore class and retrieve() below for how it's used.
# Imported defensively like the format libraries above: a clear error
# only appears if something tries to actually build or load an index
# without the package installed, rather than the whole app failing to
# start.
try:
    import faiss
except ImportError:
    faiss = None

# Reads the .env file in this folder (if present) and copies any
# KEY=VALUE lines into the environment - so GEMINI_API_KEY doesn't
# need to be retyped in every new terminal session. Safe to call even
# if no .env file exists (it just does nothing in that case).
load_dotenv()

# -------------------------------------------------------------
# CONFIGURATION DEFAULTS
# These are just the starting values. app.py lets a user override
# chunk_size / chunk_overlap / top_k / min_relevance live with
# sliders - this file defines what each knob actually does.
# See README.md -> "Tuning guide" for the full explanation of each.
# -------------------------------------------------------------
DOCS_FOLDER = "documents"
FAISS_INDEX_FILE = "faiss_index.bin"
CHUNK_METADATA_FILE = "chunk_metadata.json"
# HNSW parameters - see build_vector_store() for what each one trades off.
HNSW_M = 32              # neighbors per graph node - higher = better recall, more memory
HNSW_EF_CONSTRUCTION = 200  # build-time search depth - higher = better graph, slower to build
HNSW_EF_SEARCH = 64       # query-time search depth - higher = better recall, slower per query
HNSW_SEARCH_CANDIDATES = 200  # how many raw candidates retrieve() asks FAISS for
                                # before Python-side min_relevance/folder filtering narrows
                                # it down to top_k. Deliberately a FIXED number, NOT tied to
                                # collection size - see retrieve() for why that distinction
                                # is the entire point of using HNSW at all.
DEFAULT_MAX_EMBED_WORKERS = 5  # concurrent embedding calls during indexing - see build_vector_store()
EMBEDDING_MODEL = "gemini-embedding-001"
EMBEDDING_DIM = 768          # Gemini embeddings default to 3072 numbers per
                              # vector; we ask for a smaller 768-number version
                              # instead - see README "Tuning guide" for why.
CHAT_MODEL = "gemini-2.5-flash"
# If CHAT_MODEL is temporarily overloaded (a 503 error - see generate_answer
# below), these are tried next, in order. All three are separately free,
# so a demand spike on one doesn't take the whole app down with it.
CHAT_MODEL_FALLBACKS = ["gemini-2.5-flash-lite", "gemini-3.5-flash-lite"]
DEFAULT_CHUNK_SIZE = 500
DEFAULT_CHUNK_OVERLAP = 50
DEFAULT_TOP_K = 3
DEFAULT_MIN_RELEVANCE = 0.15

# genai.Client() reads your API key automatically from the
# GEMINI_API_KEY environment variable - see README.md "Setup".
client = genai.Client()


# ================================================================
# OFFLINE PIPELINE: Documents -> Chunking -> Embedding -> Vector DB
# ================================================================

def extract_text_from_txt(filepath):
    with open(filepath, "r", encoding="utf-8") as f:
        return f.read()


def extract_text_from_pdf(filepath):
    """
    Extracts text page by page with pdfplumber, inserting a "[Page N]"
    marker before each page's text. If a chunk later spans or starts
    near a page boundary, that marker stays visible right inside the
    chunk text - a simple way to keep page context without needing a
    more complex position-tracking system.

    Note: this only extracts the text LAYER of a PDF. A scanned PDF
    (a photo of a page with no real text underneath) will come back
    empty - that's a known limitation, flagged in the README.
    """
    if pdfplumber is None:
        raise ImportError("PDF support needs the 'pdfplumber' package. Run: pip install pdfplumber")

    pages_text = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text() or ""
            if page_text.strip():
                pages_text.append(f"[Page {i}]\n{page_text.strip()}")
    return "\n\n".join(pages_text)


def extract_text_from_docx(filepath):
    """
    Extracts paragraph text in reading order, plus any tables -
    flattened into readable "Header: value | Header: value" style
    rows so table data (like a benefits cost table) stays searchable
    even though the original table structure is lost in the process.
    """
    if DocxDocument is None:
        raise ImportError("Word support needs the 'python-docx' package. Run: pip install python-docx")

    doc = DocxDocument(filepath)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_xlsx(filepath):
    """
    Extracts every sheet's rows as readable "Header: value" lines,
    using row 1 as column headers when present. This turns a
    spreadsheet into self-contained, searchable sentences (e.g.
    "Vendor Name: Acme | Category: Office Supplies | ...") instead of
    the raw grid a keyword search would otherwise have to guess at.
    """
    if openpyxl is None:
        raise ImportError("Excel support needs the 'openpyxl' package. Run: pip install openpyxl")

    wb = openpyxl.load_workbook(filepath, data_only=True)
    parts = []
    for sheet in wb.worksheets:
        rows = list(sheet.iter_rows(values_only=True))
        if not rows:
            continue
        parts.append(f"[Sheet: {sheet.title}]")
        header = [str(c) if c is not None else "" for c in rows[0]]
        has_header = any(h.strip() for h in header)
        data_rows = rows[1:] if has_header else rows
        for row in data_rows:
            if row is None or all(c is None for c in row):
                continue
            if has_header:
                line = " | ".join(
                    f"{header[i].strip()}: {row[i]}" for i in range(len(row))
                    if i < len(header) and header[i].strip() and row[i] is not None
                )
            else:
                line = " | ".join(str(c) for c in row if c is not None)
            if line.strip():
                parts.append(line)
    return "\n".join(parts)


class _HTMLTextExtractor(HTMLParser):
    """Minimal HTML-to-text converter: keeps only the text nodes,
    drops every tag. Used when an email has no plain-text part, only
    HTML - common for HR/marketing-tool-sent emails."""
    def __init__(self):
        super().__init__()
        self._parts = []

    def handle_data(self, data):
        self._parts.append(data)

    def get_text(self):
        text = "".join(self._parts)
        # collapse the blank-line noise HTML->text tends to leave behind
        lines = [line.strip() for line in text.splitlines()]
        return "\n".join(line for line in lines if line)


def extract_text_from_eml(filepath):
    """
    Extracts sender/recipient/subject/date headers plus the message
    body from a standard .eml file (the portable, universal email
    format - every email client, including Outlook, can export or
    save messages this way). Uses Python's own built-in `email`
    module - no extra package needed.

    Prefers the plain-text body if the email has one; falls back to
    stripping tags from the HTML body if it doesn't (some emails,
    especially from HR/marketing tools, are HTML-only).
    """
    with open(filepath, "rb") as f:
        msg = _EmailBytesParser(policy=_email_policy.default).parse(f)

    headers = (
        f"From: {msg.get('From', '(unknown)')}\n"
        f"To: {msg.get('To', '(unknown)')}\n"
        f"Subject: {msg.get('Subject', '(no subject)')}\n"
        f"Date: {msg.get('Date', '(unknown)')}\n\n"
    )

    body_part = msg.get_body(preferencelist=("plain", "html"))
    if body_part is None:
        body_text = ""
    elif body_part.get_content_type() == "text/html":
        body_text = _html_to_text_helper(body_part.get_content())
    else:
        body_text = body_part.get_content()

    return headers + body_text.strip()


def _html_to_text_helper(html_content):
    """Small wrapper so extract_text_from_eml reads cleanly above -
    HTMLParser.feed() returns None, so the actual text has to be
    pulled from the parser instance afterward, not from feed()'s
    return value."""
    parser = _HTMLTextExtractor()
    parser.feed(html_content)
    return parser.get_text()


def extract_text_from_msg(filepath):
    """
    Extracts the same fields as extract_text_from_eml, but from
    Outlook's proprietary binary .msg format instead. Needs the
    third-party 'extract-msg' package (pure Python, cross-platform -
    not the same as win32com, which only works on Windows with
    Outlook actually installed).

    Note in the README: this path is built carefully against
    extract-msg's documented API, but - unlike every other extractor
    in this file - it was NOT verified against a real .msg file
    before shipping, since generating a realistic one requires
    Outlook itself. Test this one yourself with a real exported email
    before relying on it.
    """
    if extract_msg is None:
        raise ImportError("Outlook .msg support needs the 'extract-msg' package. Run: pip install extract-msg")

    msg = extract_msg.Message(filepath)
    try:
        headers = (
            f"From: {msg.sender or '(unknown)'}\n"
            f"To: {msg.to or '(unknown)'}\n"
            f"Subject: {msg.subject or '(no subject)'}\n"
            f"Date: {msg.date or '(unknown)'}\n\n"
        )
        body = msg.body or ""
        return headers + body.strip()
    finally:
        msg.close()


def extract_text_from_db(filepath):
    """
    Extracts every table in a SQLite database file as readable
    "Column: value" lines, one per row - the same idea as the Excel
    extractor above, just for a database file instead of a
    spreadsheet. Uses Python's own built-in `sqlite3` module - no
    extra package needed.

    SCOPE NOTE, worth knowing: this reads a SQLite FILE sitting in
    documents/, the same way every other format here works - it does
    NOT connect to a live database server (Postgres, MySQL, SQL
    Server). Supporting a live server would mean handling connection
    strings, credentials, and network access - a meaningfully
    different and more sensitive feature than "read this file." A
    real company using a server database would export or snapshot
    the relevant tables to a .db file periodically and drop that file
    here, the same way they'd drop in an exported spreadsheet.
    """
    conn = _sqlite3_db.connect(filepath)
    conn.row_factory = _sqlite3_db.Row
    try:
        tables = [row[0] for row in conn.execute(
            "SELECT name FROM sqlite_master WHERE type='table' AND name NOT LIKE 'sqlite_%'"
        ).fetchall()]

        parts = []
        for table in tables:
            rows = conn.execute(f"SELECT * FROM [{table}]").fetchall()
            if not rows:
                continue
            parts.append(f"[Table: {table}]")
            columns = rows[0].keys()
            for row in rows:
                line = " | ".join(
                    f"{col}: {row[col]}" for col in columns if row[col] is not None
                )
                if line.strip():
                    parts.append(line)
        return "\n".join(parts)
    finally:
        conn.close()


# Maps a file extension to the function that turns that file into
# plain text. Add a new format by writing an extract_text_from_X()
# function above and adding one line here - load_documents() below
# doesn't need to change at all.
EXTRACTORS = {
    ".txt": extract_text_from_txt,
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".xlsx": extract_text_from_xlsx,
    ".eml": extract_text_from_eml,
    ".msg": extract_text_from_msg,
    ".db": extract_text_from_db,
    ".sqlite": extract_text_from_db,
    ".sqlite3": extract_text_from_db,
}


def load_documents(root_folder=DOCS_FOLDER, skipped_log=None):
    """
    Walk through every subfolder of `root_folder` and extract text
    from every file whose extension is in EXTRACTORS (currently .txt,
    .pdf, .docx, .xlsx). This is what makes folder-level citations
    possible: documents/HR_Policies/leave_policy.txt keeps
    "HR_Policies" attached, so later we can tell the user not just
    WHICH file a fact came from, but WHERE that file lives.

    Any file that fails to extract (corrupt file, unsupported format,
    a required package not installed) is SKIPPED rather than crashing
    the whole indexing run - if `skipped_log` is given (a list),
    a short reason is appended to it so the caller can surface what
    got skipped and why, instead of a document silently vanishing.

    Returns a list of dicts: {folder, filename, path, text}
    """
    documents = []
    pattern = os.path.join(root_folder, "**", "*")
    for filepath in glob.glob(pattern, recursive=True):
        if not os.path.isfile(filepath):
            continue

        ext = os.path.splitext(filepath)[1].lower()
        extractor = EXTRACTORS.get(ext)
        if extractor is None:
            continue  # not a supported format - skip silently (e.g. .DS_Store)

        try:
            text = extractor(filepath)
        except Exception as e:
            if skipped_log is not None:
                skipped_log.append(f"{filepath}: {e}")
            continue

        if not text or not text.strip():
            if skipped_log is not None:
                skipped_log.append(f"{filepath}: no extractable text (empty or scanned/image-only file)")
            continue

        rel_path = os.path.relpath(filepath, root_folder)   # e.g. "HR_Policies/leave_policy.txt"
        folder = os.path.dirname(rel_path) or "(root)"       # e.g. "HR_Policies"
        filename = os.path.basename(rel_path)                 # e.g. "leave_policy.txt"
        documents.append({
            "folder": folder,
            "filename": filename,
            "path": rel_path,     # <- the "ref + folder" citation is built from this
            "text": text,
            # SHA-256 of the EXTRACTED TEXT, not the raw file bytes -
            # this matters because a document can be re-saved by its
            # source application (Word, Excel...) with a totally
            # different internal byte layout while the actual
            # readable content is identical. Hashing the text means
            # build_vector_store() only re-embeds a file when what a
            # person would actually read has changed, not whenever a
            # file's bytes happen to differ.
            "text_hash": hashlib.sha256(text.encode("utf-8")).hexdigest(),
        })
    return documents


def chunk_text(text, chunk_size=DEFAULT_CHUNK_SIZE, overlap=DEFAULT_CHUNK_OVERLAP):
    """
    Split one long string into overlapping pieces of `chunk_size`
    characters each.

    WHY CHUNK: a vector search step needs small, focused pieces of
    text - if we embedded an entire document as one vector, a
    question about paragraph 4 would get diluted by paragraphs 1-3
    and 5-10 all being mashed into the same vector.

    WHY OVERLAP: without it, a sentence sitting exactly on a chunk
    boundary gets cut in half and neither half makes sense alone.
    The overlap re-includes the last `overlap` characters of the
    previous chunk at the start of the next one.

    Example, chunk_size=10, overlap=3, text="ABCDEFGHIJKLMNOPQRSTUVWXYZ":
        chunk 0: "ABCDEFGHIJ"
        chunk 1: "HIJKLMNOPQ"   (starts at 10-3=7, re-using H,I,J)
        chunk 2: "OPQRSTUVWX"
        chunk 3: "VWXYZ"
    (Verified with a standalone test before this file was written.)
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        piece = text[start:end]
        if piece.strip():           # don't keep a chunk that's just whitespace
            chunks.append(piece)
        start = end - overlap
        if overlap >= chunk_size:   # safety net: this combination would loop forever otherwise
            break
    return chunks


def _is_transient_error(error):
    """
    Two very different things can go wrong when calling Gemini:

    1. A 503 "UNAVAILABLE / high demand" error - Google's servers are
       temporarily overloaded. This affects every user, on every plan,
       free or paid - it has nothing to do with your code or your key.
       It almost always succeeds a few seconds later.
    2. A 429 "RESOURCE_EXHAUSTED" error - YOU personally have used up
       your free-tier quota for this minute/day. Also usually temporary.

    Both are worth retrying automatically. Anything else (like a bad
    API key, or a malformed request) will NOT fix itself by retrying,
    so we let those raise immediately instead of silently hiding a
    real bug behind a retry loop.
    """
    message = str(error)
    return "503" in message or "UNAVAILABLE" in message or "429" in message or "RESOURCE_EXHAUSTED" in message


def _call_with_retry(api_call, max_attempts=3):
    """
    Calls `api_call()` and, if it fails with a transient error, waits
    and tries again - up to `max_attempts` times. The wait time doubles
    each attempt (1s, 2s, 4s...) with a little random jitter added, so
    that if many people hit the same overloaded model at once, they
    don't all retry at exactly the same moment and cause a second wave
    of overload (this pattern is called "exponential backoff").
    """
    last_error = None
    for attempt in range(max_attempts):
        try:
            return api_call()
        except Exception as e:
            if not _is_transient_error(e) or attempt == max_attempts - 1:
                raise
            last_error = e
            wait_seconds = (2 ** attempt) + random.uniform(0, 1)
            time.sleep(wait_seconds)
    raise last_error


def embed_text(text):
    """
    Turn one piece of text into a vector: a list of 768 numbers that
    represents its MEANING. Two chunks that mean similar things
    produce vectors that point in a similar direction - that's the
    entire trick behind "semantic search" (matching meaning, not
    keywords). This is the one function that talks to Google's
    servers - everything else here is local computation. Wrapped in
    _call_with_retry so a momentary server hiccup doesn't crash the
    whole indexing run.
    """
    def _call():
        return client.models.embed_content(
            model=EMBEDDING_MODEL,
            contents=text,
            config=types.EmbedContentConfig(output_dimensionality=EMBEDDING_DIM)
        )
    result = _call_with_retry(_call)
    return result.embeddings[0].values


class VectorStore:
    """
    Wraps a FAISS index (the actual similarity-search engine) together
    with the metadata for each chunk - folder, filename, path, chunk
    text, etc. FAISS itself only knows about numbers; it has no idea
    what a "folder" or a "citation" is. `metadata[i]` always describes
    the same chunk as vector `i` inside `index` - they're kept in
    sync by always appending to both together, in the same order, in
    build_vector_store() below.

    __len__ and __iter__ deliberately make this behave like the plain
    list of chunk dicts the app used before FAISS existed - so
    app.py's existing `len(vector_store)` and
    `for entry in vector_store` code keeps working completely
    unchanged. That's on purpose: swapping the search engine
    underneath shouldn't require touching the UI at all.
    """
    def __init__(self, index, metadata):
        self.index = index
        self.metadata = metadata

    def __len__(self):
        return len(self.metadata)

    def __iter__(self):
        return iter(self.metadata)


def normalize_vector(v):
    """
    L2-normalizes a vector (scales it to length 1, keeping its
    direction). This matters because FAISS's inner-product metric
    computes a raw dot product, not cosine similarity - but the dot
    product of two NORMALIZED vectors is mathematically exactly equal
    to their cosine similarity. Normalizing both the stored vectors
    (in build_vector_store) and every query (in retrieve) is what
    makes FAISS's scores mean exactly the same thing as the
    cosine_similarity() this project used before FAISS was
    introduced - so every existing relevance threshold (0.15, 0.35,
    etc.) still means the same thing, with no retuning needed. This
    only holds if the index was actually built with
    faiss.METRIC_INNER_PRODUCT - see build_vector_store().
    """
    arr = np.asarray(v, dtype="float32")
    norm = np.linalg.norm(arr)
    if norm == 0:
        return arr
    return arr / norm


def build_vector_store(chunk_size=DEFAULT_CHUNK_SIZE, chunk_overlap=DEFAULT_CHUNK_OVERLAP,
                        root_folder=DOCS_FOLDER, progress_callback=None,
                        previous_store=None, max_workers=DEFAULT_MAX_EMBED_WORKERS):
    """
    Runs the full offline pipeline: load -> chunk -> embed -> build a
    searchable index. Returns a VectorStore.

    THREE THINGS THIS FUNCTION DOES TO STAY FAST AT SCALE:

    1. HNSW instead of exact search. Uses faiss.IndexHNSWFlat, not
       IndexFlatIP - a real algorithmic change, not just a faster
       implementation of the same algorithm. IndexFlatIP compares a
       query against EVERY stored vector, always - correct, but
       inherently linear in collection size. HNSW builds a navigable
       graph over the vectors at index time (governed by HNSW_M and
       HNSW_EF_CONSTRUCTION above) so that at query time, a search
       only has to visit a small, mostly-constant-size neighborhood
       of that graph (governed by HNSW_EF_SEARCH), not the whole
       collection. That's what makes search time stop growing
       linearly with the number of documents. The honest tradeoff:
       HNSW is APPROXIMATE - with these parameters, it can very
       occasionally miss the true single best match in exchange for
       not scanning everything. At this project's scale (tens of
       chunks) that's effectively never noticeable; the values chosen
       here (M=32, efConstruction=200, efSearch=64) are solid,
       widely-used defaults that hold up well into the millions of
       vectors, not numbers picked to look good on a small demo.

    2. Parallel embedding. Each embed_text() call is a network round
       trip to Gemini - waiting for one to finish before starting the
       next means indexing time scales directly with document count.
       `max_workers` chunks are embedded CONCURRENTLY via
       ThreadPoolExecutor instead, so wall-clock indexing time is
       roughly (total chunks / max_workers) round trips, not
       (total chunks) round trips. Higher isn't free - too many
       concurrent requests can itself trigger the free tier's rate
       limit, which is why this defaults to a modest 5, not 50.

    3. Incremental re-indexing. If `previous_store` is given (the
       VectorStore from the last build), any file whose extracted
       text is byte-for-byte identical to last time (compared via
       the text_hash from load_documents()) is REUSED as-is - its
       existing chunks and embeddings are copied over, and it never
       touches the embedding API at all. Only new or actually-changed
       files get (re-)embedded. Pass previous_store=None (the
       default) to force a full rebuild from scratch, e.g. after
       changing chunk_size/chunk_overlap - a chunk boundary shift
       invalidates old chunks even if the source text didn't change,
       so reuse is deliberately skipped in that case (see app.py,
       which only passes previous_store when chunk settings are
       unchanged from what's already indexed).

    `progress_callback`, if provided, is called with a short status
    string per document (reused, freshly indexed, or skipped) - app.py
    uses this to show a live indexing log in the sidebar.
    """
    if faiss is None:
        raise ImportError("Vector search needs the 'faiss-cpu' package. Run: pip install faiss-cpu")

    skipped = []
    documents = load_documents(root_folder, skipped_log=skipped)

    # Index what was already embedded last time, by path, so an
    # unchanged file's chunks can be reused instead of re-embedded.
    # Tracks each entry's ORIGINAL position in previous_store.index too
    # - the metadata dict itself never stores the embedding (see the
    # "vectors = ..." step below, which pops it out once the FAISS
    # index is built, so a vector is never stored twice) - reusing a
    # chunk means pulling its vector back out of the previous index by
    # that original position, not finding it sitting in the metadata.
    reusable_by_path = {}
    if previous_store is not None:
        for i, entry in enumerate(previous_store.metadata):
            reusable_by_path.setdefault(entry["path"], []).append((i, entry))

    metadata = []          # final, in-order list - what actually gets returned
    pending_embeds = []    # (chunk_text, partial_metadata) pairs still needing an embedding call

    for doc in documents:
        existing = reusable_by_path.get(doc["path"])
        if existing and existing[0][1].get("text_hash") == doc["text_hash"]:
            try:
                # Pull each reused chunk's actual vector back out of the
                # PREVIOUS index by its original position - this can only
                # fail if a FAISS version genuinely can't reconstruct from
                # an HNSWFlat index, which shouldn't happen (HNSWFlat
                # stores full, un-quantized vectors by design), but if it
                # ever does, falling through to a normal re-embed keeps
                # this correct instead of crashing the whole run.
                reused_entries = []
                for original_idx, entry in existing:
                    vector = previous_store.index.reconstruct(original_idx)
                    reused_entry = dict(entry)
                    reused_entry["embedding"] = vector
                    reused_entries.append(reused_entry)
                metadata.extend(reused_entries)
                if progress_callback:
                    progress_callback(f"Reused {doc['path']} -> {len(reused_entries)} chunk(s), unchanged since last index")
                continue
            except Exception as e:
                if progress_callback:
                    progress_callback(f"Couldn't reuse {doc['path']} ({e}) - re-embedding it instead")
                # falls through to the normal chunk+embed path below

        chunks = chunk_text(doc["text"], chunk_size, chunk_overlap)
        for i, chunk in enumerate(chunks):
            pending_embeds.append((chunk, {
                "folder": doc["folder"],
                "filename": doc["filename"],
                "path": doc["path"],
                "chunk_index": i,
                "text": chunk,
                "text_hash": doc["text_hash"],
            }))
        if progress_callback:
            verb = "Indexing" if doc["path"] not in reusable_by_path else "Re-indexing (changed)"
            progress_callback(f"{verb} {doc['path']} -> {len(chunks)} chunk(s) queued for embedding")

    if pending_embeds:
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_meta = {
                executor.submit(embed_text, chunk_text_): meta
                for chunk_text_, meta in pending_embeds
            }
            for future in as_completed(future_to_meta):
                meta = dict(future_to_meta[future])
                meta["embedding"] = future.result()
                metadata.append(meta)

    if progress_callback:
        for reason in skipped:
            progress_callback(f"⚠️ Skipped {reason}")

    # Sort so the persisted metadata file reads in a sensible,
    # deterministic order regardless of parallel completion order or
    # how much was reused vs freshly embedded this run.
    metadata.sort(key=lambda m: (m["path"], m["chunk_index"]))

    vectors = [normalize_vector(m.pop("embedding")) for m in metadata]

    dim = len(vectors[0]) if vectors else EMBEDDING_DIM
    # METRIC_INNER_PRODUCT is NOT the default for IndexHNSWFlat - omitting
    # it silently gives you L2 (Euclidean) distance instead, which runs in
    # the OPPOSITE direction from cosine similarity (lower = more similar,
    # not higher) on a completely different numeric scale. Every
    # min_relevance threshold in this app (0.15, 0.35...) would silently
    # mean the opposite of what it's supposed to without this argument.
    index = faiss.IndexHNSWFlat(dim, HNSW_M, faiss.METRIC_INNER_PRODUCT)
    index.hnsw.efConstruction = HNSW_EF_CONSTRUCTION
    index.hnsw.efSearch = HNSW_EF_SEARCH
    if vectors:
        index.add(np.array(vectors, dtype="float32"))

    return VectorStore(index, metadata)


def save_vector_store(vector_store, index_file=FAISS_INDEX_FILE, metadata_file=CHUNK_METADATA_FILE):
    """Saves the FAISS index and its matching metadata as two separate
    files, so 2_ask.py doesn't have to re-embed everything (and
    re-call the API) on every run. The index is FAISS's own binary
    format (faiss.write_index) - not something to open by hand."""
    if faiss is None:
        raise ImportError("Vector search needs the 'faiss-cpu' package. Run: pip install faiss-cpu")
    faiss.write_index(vector_store.index, index_file)
    with open(metadata_file, "w", encoding="utf-8") as f:
        json.dump(vector_store.metadata, f)


def load_vector_store(index_file=FAISS_INDEX_FILE, metadata_file=CHUNK_METADATA_FILE):
    """Loads a previously-saved index back into memory - both for
    normal startup (1_ingest.py / 2_ask.py) and as the `previous_store`
    passed into build_vector_store() for incremental re-indexing."""
    if faiss is None:
        raise ImportError("Vector search needs the 'faiss-cpu' package. Run: pip install faiss-cpu")
    index = faiss.read_index(index_file)
    # Explicitly re-set efSearch after loading rather than trusting it
    # survived the save/load round trip - cheap insurance, and it
    # guarantees a loaded index always searches with the current
    # HNSW_EF_SEARCH constant even if that constant changes between
    # app versions.
    if hasattr(index, "hnsw"):
        index.hnsw.efSearch = HNSW_EF_SEARCH
    with open(metadata_file, "r", encoding="utf-8") as f:
        metadata = json.load(f)
    return VectorStore(index, metadata)


# ================================================================
# ONLINE PIPELINE: Question -> Embed & search -> Top chunks -> LLM
# ================================================================

def retrieve(question, vector_store, top_k=DEFAULT_TOP_K, min_relevance=DEFAULT_MIN_RELEVANCE, allowed_folders=None):
    """
    Embed the question, ask FAISS to rank the closest chunks against
    it, throw out anything below `min_relevance` OR outside
    `allowed_folders`, and return the `top_k` highest-scoring chunks,
    best first.

    `allowed_folders`: None or "ALL" means no restriction (an admin,
    or code that hasn't been updated to pass this at all - e.g.
    1_ingest.py/2_ask.py, which have no login system and keep working
    exactly as before). Otherwise, a list/set of folder names - any
    chunk whose folder isn't in that set is filtered out HERE, before
    it's ever assembled into a prompt or shown to anyone. This is
    what makes department-level access control real enforcement
    rather than something the UI merely chooses not to display -
    a restricted chunk is never retrieved in the first place, so it
    can't leak through a different code path later by accident.

    Returns a list of (score, chunk_record) tuples - same shape as
    before FAISS was introduced, so build_prompt() and everything
    downstream needs no changes at all.

    Asks FAISS for HNSW_SEARCH_CANDIDATES results (a fixed number,
    NOT the full collection) before filtering by min_relevance and
    allowed_folders - this is deliberately bounded, and the bound is
    the entire reason HNSW is fast. Asking an approximate index for
    "every vector, ranked" forces it to behave like an exhaustive
    scan again, which would silently throw away the whole point of
    using HNSW over IndexFlatIP in the first place.

    Honest tradeoff worth knowing: because filtering happens AFTER a
    bounded fetch, a chunk that's genuinely relevant but ranks outside
    the top HNSW_SEARCH_CANDIDATES globally - realistic only in a
    collection with far more documents than this project has, for a
    user restricted to a small slice of it - could be missed even
    though it exists in an allowed folder. HNSW_SEARCH_CANDIDATES=200
    comfortably covers this project's real scale with room to spare.
    The genuine production fix at massive multi-department scale is
    FAISS's own metadata pre-filtering (an IDSelector restricting
    which vectors are even considered during the search itself,
    rather than filtering results after) - a real, known next step,
    not implemented here.
    """
    if len(vector_store) == 0:
        return []

    restrict_folders = allowed_folders is not None and allowed_folders != "ALL"
    if restrict_folders:
        allowed_folders = set(allowed_folders)

    question_vector = normalize_vector(embed_text(question))
    query_array = np.array([question_vector], dtype="float32")

    search_k = min(len(vector_store), HNSW_SEARCH_CANDIDATES)
    scores, indices = vector_store.index.search(query_array, search_k)

    results = []
    for score, idx in zip(scores[0], indices[0]):
        if idx == -1:  # FAISS pads with -1 if fewer than search_k vectors exist
            continue
        if score < min_relevance:
            continue
        entry = vector_store.metadata[idx]
        if restrict_folders and entry["folder"] not in allowed_folders:
            continue
        results.append((float(score), entry))

    # FAISS already returns results ordered best-match-first for an
    # inner-product index, so no re-sort is needed here.
    return results[:top_k]


def build_prompt(question, top_chunks, chat_history=None, max_history_turns=4):
    """
    Assemble the system + user prompt sent to the LLM. Each
    retrieved chunk is numbered [1], [2]... and labeled with its
    full folder/file path so the model can cite exactly where a
    fact came from - this is what makes "[1] HR_Policies/leave_
    policy.txt (chunk 0)" show up in the final answer.

    `chat_history`, if given, is a list of {"role": "user"/"assistant",
    "content": ...} dicts - the last `max_history_turns` of them are
    included so the LLM can answer in a natural conversational way
    ("as mentioned above...") instead of treating every message as
    if it's the first one. The LLM still must ground every fact in
    the numbered context below, not in what it remembers saying.
    """
    context_block = ""
    for i, (score, entry) in enumerate(top_chunks, start=1):
        context_block += f"[{i}] {entry['path']} (chunk {entry['chunk_index']})\n{entry['text']}\n\n"

    history_block = ""
    if chat_history:
        recent = chat_history[-max_history_turns:]
        history_lines = "\n".join(f"{turn['role']}: {turn['content']}" for turn in recent)
        history_block = f"Recent conversation (for tone/continuity only):\n{history_lines}\n\n"

    system_prompt = (
        "You are an internal company document assistant having a conversation "
        "with an employee. Answer ONLY using the numbered context below - never "
        "the recent conversation - for facts. "
        "If the answer is not in the context, respond with a short plain sentence "
        "saying you don't know - do NOT include any bracket citations like [1] in "
        "that sentence, since you are not citing a fact, only declining to answer. "
        "Never invent policy details. "
        "If you DO answer from the context, cite the source after every factual "
        "claim in brackets, like [1], matching the numbered context blocks given "
        "to you. You may refer naturally to earlier parts of the conversation for "
        "tone and flow, but never as a substitute for a citation."
    )
    user_prompt = f"{history_block}Context:\n{context_block}\nQuestion: {question}"
    return system_prompt, user_prompt


def generate_answer(question, top_chunks, chat_history=None, temperature=0):
    """
    Send the assembled prompt to the chat model and return its answer.
    temperature=0 means "be as deterministic and literal as possible" -
    the right setting for a factual company assistant, where you do
    NOT want creative variation between two people asking the same
    question.

    Tries CHAT_MODEL first, with a few retries if Google's servers are
    briefly overloaded (see _call_with_retry). If that model keeps
    failing after all retries, moves on to each model in
    CHAT_MODEL_FALLBACKS in turn, rather than giving up immediately -
    this is what stops a single overloaded model from crashing a live
    demo.
    """
    system_prompt, user_prompt = build_prompt(question, top_chunks, chat_history=chat_history)
    config = types.GenerateContentConfig(
        system_instruction=system_prompt,
        temperature=temperature
    )

    models_to_try = [CHAT_MODEL] + CHAT_MODEL_FALLBACKS
    last_error = None

    for model_name in models_to_try:
        def _call(model_name=model_name):
            return client.models.generate_content(
                model=model_name,
                contents=user_prompt,
                config=config
            )
        try:
            response = _call_with_retry(_call, max_attempts=2)
            return response.text
        except Exception as e:
            last_error = e
            continue  # this model is out - try the next one in the list

    # Every model in the list failed - re-raise the last error so
    # app.py can show the user a clear, friendly message instead of
    # a raw crash.
    raise last_error


def condense_question(chat_history, new_question, max_history_turns=4):
    """
    Rewrites a possibly context-dependent follow-up ("what's the warranty
    on it?") into a fully standalone question ("what's the warranty on
    the X200 pump?") using recent conversation turns. This standalone
    version is what actually gets embedded and searched in retrieve() -
    without this step, follow-up questions would search for the literal
    words "it" and "warranty" with no idea what "it" refers to, and
    retrieval quality would collapse after the first message.

    If there's no history yet (first message in the chat), the question
    is already standalone, so this returns it unchanged without spending
    an extra API call.
    """
    if not chat_history:
        return new_question

    recent = chat_history[-max_history_turns:]
    history_lines = "\n".join(f"{turn['role']}: {turn['content']}" for turn in recent)

    system_prompt = (
        "Rewrite the user's latest message as a fully standalone question "
        "that makes sense with no prior context. Do not answer it - only "
        "rewrite it. If it is already standalone, return it unchanged. "
        "Return ONLY the rewritten question, nothing else."
    )
    user_prompt = f"Conversation so far:\n{history_lines}\n\nLatest message: {new_question}\n\nStandalone question:"
    config = types.GenerateContentConfig(system_instruction=system_prompt, temperature=0)

    for model_name in [CHAT_MODEL] + CHAT_MODEL_FALLBACKS:
        def _call(model_name=model_name):
            return client.models.generate_content(model=model_name, contents=user_prompt, config=config)
        try:
            response = _call_with_retry(_call, max_attempts=2)
            return response.text.strip()
        except Exception:
            continue  # condensing failed - fall through to the raw question below

    return new_question  # never crash the turn just because condensing failed


def answer_question(question, vector_store, top_k=DEFAULT_TOP_K, min_relevance=DEFAULT_MIN_RELEVANCE, allowed_folders=None):
    """
    Convenience function used by 1_ingest.py/2_ask.py: runs the single-
    turn online pipeline (retrieve -> generate) with no conversation
    memory. Returns (answer_text, top_chunks) so the caller can show
    sources alongside the answer.

    `allowed_folders` defaults to None (no restriction) since the CLI
    tools have no login system - this parameter only matters once a
    caller (app.py) actually has a logged-in user to restrict.
    """
    top_chunks = retrieve(question, vector_store, top_k, min_relevance, allowed_folders=allowed_folders)
    if not top_chunks:
        return (
            "I couldn't find anything in the indexed documents that answers "
            "this with enough confidence. Try rephrasing the question, or "
            "lower the minimum relevance score in the sidebar.",
            []
        )
    answer = generate_answer(question, top_chunks)
    return answer, top_chunks


def answer_question_conversational(question, vector_store, chat_history=None,
                                     top_k=DEFAULT_TOP_K, min_relevance=DEFAULT_MIN_RELEVANCE,
                                     allowed_folders=None):
    """
    Full conversational RAG turn, used by app.py's chat UI:
      1. condense_question()  - rewrite a possibly ambiguous follow-up
         into a standalone question, using recent chat history
      2. retrieve()           - search using the standalone question,
         filtered to `allowed_folders` if the caller passes one
      3. generate_answer()    - answer using the retrieved chunks AND
         the raw chat history (for natural tone), citing sources as usual

    Returns (answer_text, top_chunks, standalone_question) - the third
    value is useful to show in the UI so the user can see exactly what
    was searched for, which is good for both debugging and demoing that
    follow-up questions are genuinely being understood in context.
    """
    standalone_question = condense_question(chat_history, question)
    top_chunks = retrieve(standalone_question, vector_store, top_k, min_relevance, allowed_folders=allowed_folders)

    if not top_chunks:
        return (
            "I couldn't find anything in the indexed documents that answers "
            "this with enough confidence. Try rephrasing, or use a broader "
            "search mode.",
            [],
            standalone_question
        )

    answer = generate_answer(question, top_chunks, chat_history=chat_history)
    return answer, top_chunks, standalone_question